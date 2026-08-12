"""待审正文解析器 — 支持 PDF / Word / TXT / 纯文本输入."""

import re
from io import BytesIO
from pathlib import Path
from typing import Any, Optional, Union

from geo_review.models import ParsedContent


_MAX_CONTENT_LENGTH = 500000
_MIN_TEXT_LENGTH = 10


def _clean_text(text: str) -> str:
    """清洗提取的文本：合并空白、去除多余换行等."""
    if not text:
        return ""

    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")

    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)

    text = text.strip()
    return text


def _truncate_text(text: str, max_len: int = _MAX_CONTENT_LENGTH) -> tuple[str, bool]:
    """截断文本并标记是否被截断."""
    if len(text) <= max_len:
        return text, False
    truncated = text[:max_len].rsplit("\n", 1)[0]
    return truncated, True


class ContentParser:
    """待审正文解析器 — 统一入口.

    支持输入:
        - 纯文本字符串
        - PDF 文件 (.pdf)
        - Word 文件 (.docx / .doc)
        - TXT 文件 (.txt)
        - 文件二进制内容 (bytes)

    返回:
        - ParsedContent 对象（包含提取的文本、来源信息、警告等）
    """

    # ------------------------------------------------------------------
    # 公共入口
    # ------------------------------------------------------------------
    @classmethod
    def parse(
        cls,
        source: Union[str, Path, bytes],
        *,
        filename: Optional[str] = None,
        format_hint: Optional[str] = None,
        max_length: int = _MAX_CONTENT_LENGTH,
    ) -> ParsedContent:
        """自动识别并解析待审正文.

        Args:
            source: 文件路径、纯文本内容或文件二进制数据
            filename: 可选，用于辅助判断格式（当 source 为 bytes 时）
            format_hint: 强制指定格式 ('pdf' | 'docx' | 'doc' | 'txt' | 'text')
            max_length: 内容最大字符数限制

        Returns:
            ParsedContent 对象

        Raises:
            ValueError: 解析失败或文本内容过短
        """
        fmt = (format_hint or "").lower()

        if isinstance(source, str):
            stripped = source.strip()
            if not stripped:
                raise ValueError("正文内容为空")

            if stripped.startswith(("{", "[")):
                raise ValueError("JSON 格式不是合法的正文输入")

            if fmt == "text":
                return cls._from_text(stripped, max_length=max_length)

            if fmt:
                if fmt in ("pdf", "docx", "doc", "txt"):
                    path = Path(stripped)
                    if not path.exists():
                        raise FileNotFoundError(f"文件不存在: {path}")
                    return cls._from_file(path, max_length=max_length)
                else:
                    raise ValueError(f"不支持的格式: {fmt}")

            try:
                path = Path(stripped)
                if path.exists():
                    return cls._from_file(path, max_length=max_length)
            except Exception:
                pass

            return cls._from_text(stripped, max_length=max_length)

        if isinstance(source, Path):
            return cls._from_file(source, max_length=max_length)

        if isinstance(source, bytes):
            if not source:
                raise ValueError("文件内容为空")

            detected_format = fmt or cls._detect_format(source, filename)
            return cls._from_bytes(source, detected_format, filename, max_length)

        raise TypeError(f"不支持的输入类型: {type(source)}")

    # ------------------------------------------------------------------
    # 各格式专用方法
    # ------------------------------------------------------------------
    @classmethod
    def _from_text(cls, text: str, *, max_length: int) -> ParsedContent:
        """从纯文本解析."""
        cleaned = _clean_text(text)
        if len(cleaned) < _MIN_TEXT_LENGTH:
            raise ValueError(f"正文内容过短（{len(cleaned)} 字符）")

        truncated_text, truncated = _truncate_text(cleaned, max_length)

        return ParsedContent(
            text=truncated_text,
            source="text",
            char_count=len(truncated_text),
            truncated=truncated,
        )

    @classmethod
    def _from_file(cls, path: Path, *, max_length: int) -> ParsedContent:
        """从文件解析."""
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return cls._from_pdf_file(path, max_length)
        if suffix == ".docx":
            return cls._from_docx_file(path, max_length)
        if suffix == ".doc":
            return cls._from_doc_file(path, max_length)
        if suffix == ".txt":
            return cls._from_txt_file(path, max_length)

        raise ValueError(f"不支持的文件格式: {suffix}")

    @classmethod
    def _from_bytes(
        cls,
        data: bytes,
        fmt: str,
        filename: Optional[str],
        max_length: int,
    ) -> ParsedContent:
        """从二进制数据解析."""
        if fmt == "pdf":
            return cls._from_pdf_bytes(data, filename, max_length)
        if fmt == "docx":
            return cls._from_docx_bytes(data, filename, max_length)
        if fmt == "doc":
            return cls._from_doc_bytes(data, filename, max_length)
        if fmt == "txt":
            return cls._from_txt_bytes(data, filename, max_length)

        raise ValueError(f"不支持的格式: {fmt}")

    # ------------------------------------------------------------------
    # PDF 解析
    # ------------------------------------------------------------------
    @classmethod
    def _from_pdf_file(cls, path: Path, max_length: int) -> ParsedContent:
        """从 PDF 文件提取文本."""
        try:
            from PyPDF2 import PdfReader
        except ImportError as exc:
            raise ImportError("请安装 PyPDF2: pip install PyPDF2") from exc

        warnings: list[str] = []

        try:
            reader = PdfReader(str(path))
            page_count = len(reader.pages)

            if page_count == 0:
                raise ValueError("PDF 文件为空")

            texts: list[str] = []
            for page in reader.pages:
                try:
                    text = page.extract_text() or ""
                    if text.strip():
                        texts.append(text)
                except Exception as exc:
                    warnings.append(f"解析某页失败: {exc}")

            if not texts:
                warnings.append("PDF 可能为纯图片格式，无法提取文本")
                raise ValueError("PDF 中未提取到文本内容")

            text = "\n\n".join(texts)
            cleaned = _clean_text(text)

            if len(cleaned) < _MIN_TEXT_LENGTH:
                warnings.append("提取的文本内容较短")

            truncated_text, truncated = _truncate_text(cleaned, max_length)

            return ParsedContent(
                text=truncated_text,
                source="pdf",
                filename=path.name,
                page_count=page_count,
                char_count=len(truncated_text),
                truncated=truncated,
                warnings=warnings,
            )

        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"PDF 解析失败: {exc}") from exc

    @classmethod
    def _from_pdf_bytes(cls, data: bytes, filename: Optional[str], max_length: int) -> ParsedContent:
        """从 PDF 二进制数据提取文本."""
        try:
            from PyPDF2 import PdfReader
        except ImportError as exc:
            raise ImportError("请安装 PyPDF2: pip install PyPDF2") from exc

        warnings: list[str] = []

        try:
            reader = PdfReader(BytesIO(data))
            page_count = len(reader.pages)

            if page_count == 0:
                raise ValueError("PDF 文件为空")

            texts: list[str] = []
            for page in reader.pages:
                try:
                    text = page.extract_text() or ""
                    if text.strip():
                        texts.append(text)
                except Exception as exc:
                    warnings.append(f"解析某页失败: {exc}")

            if not texts:
                warnings.append("PDF 可能为纯图片格式，无法提取文本")
                raise ValueError("PDF 中未提取到文本内容")

            text = "\n\n".join(texts)
            cleaned = _clean_text(text)

            if len(cleaned) < _MIN_TEXT_LENGTH:
                warnings.append("提取的文本内容较短")

            truncated_text, truncated = _truncate_text(cleaned, max_length)

            return ParsedContent(
                text=truncated_text,
                source="pdf",
                filename=filename,
                page_count=page_count,
                char_count=len(truncated_text),
                truncated=truncated,
                warnings=warnings,
            )

        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"PDF 解析失败: {exc}") from exc

    # ------------------------------------------------------------------
    # DOCX 解析
    # ------------------------------------------------------------------
    @classmethod
    def _from_docx_file(cls, path: Path, max_length: int) -> ParsedContent:
        """从 DOCX 文件提取文本."""
        try:
            import docx
        except ImportError as exc:
            raise ImportError("请安装 python-docx: pip install python-docx") from exc

        warnings: list[str] = []

        try:
            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            if not paragraphs:
                raise ValueError("DOCX 文件中未提取到文本内容")

            text = "\n\n".join(paragraphs)
            cleaned = _clean_text(text)

            if len(cleaned) < _MIN_TEXT_LENGTH:
                warnings.append("提取的文本内容较短")

            truncated_text, truncated = _truncate_text(cleaned, max_length)

            return ParsedContent(
                text=truncated_text,
                source="docx",
                filename=path.name,
                char_count=len(truncated_text),
                truncated=truncated,
                warnings=warnings,
            )

        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"DOCX 解析失败: {exc}") from exc

    @classmethod
    def _from_docx_bytes(cls, data: bytes, filename: Optional[str], max_length: int) -> ParsedContent:
        """从 DOCX 二进制数据提取文本."""
        try:
            import docx
        except ImportError as exc:
            raise ImportError("请安装 python-docx: pip install python-docx") from exc

        warnings: list[str] = []

        try:
            doc = docx.Document(BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            if not paragraphs:
                raise ValueError("DOCX 文件中未提取到文本内容")

            text = "\n\n".join(paragraphs)
            cleaned = _clean_text(text)

            if len(cleaned) < _MIN_TEXT_LENGTH:
                warnings.append("提取的文本内容较短")

            truncated_text, truncated = _truncate_text(cleaned, max_length)

            return ParsedContent(
                text=truncated_text,
                source="docx",
                filename=filename,
                char_count=len(truncated_text),
                truncated=truncated,
                warnings=warnings,
            )

        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"DOCX 解析失败: {exc}") from exc

    # ------------------------------------------------------------------
    # DOC 解析
    # ------------------------------------------------------------------
    @classmethod
    def _from_doc_file(cls, path: Path, max_length: int) -> ParsedContent:
        """从 DOC 文件提取文本（使用 antiword 或 docx2txt）."""
        warnings: list[str] = []

        try:
            import docx2txt
            text = docx2txt.process(str(path))
        except ImportError:
            try:
                text = cls._extract_doc_with_antiword(path)
            except Exception as fallback_exc:
                raise ImportError(
                    "请安装 docx2txt: pip install docx2txt（或安装 antiword 命令行工具）"
                ) from fallback_exc
        except Exception as exc:
            raise ValueError(f"DOC 解析失败: {exc}") from exc

        if not text or not text.strip():
            raise ValueError("DOC 文件中未提取到文本内容")

        cleaned = _clean_text(text)

        if len(cleaned) < _MIN_TEXT_LENGTH:
            warnings.append("提取的文本内容较短")

        truncated_text, truncated = _truncate_text(cleaned, max_length)

        return ParsedContent(
            text=truncated_text,
            source="doc",
            filename=path.name,
            char_count=len(truncated_text),
            truncated=truncated,
            warnings=warnings,
        )

    @classmethod
    def _from_doc_bytes(cls, data: bytes, filename: Optional[str], max_length: int) -> ParsedContent:
        """从 DOC 二进制数据提取文本."""
        warnings: list[str] = []

        try:
            import docx2txt
            text = docx2txt.process(BytesIO(data))
        except ImportError:
            raise ImportError("请安装 docx2txt: pip install docx2txt")
        except Exception as exc:
            raise ValueError(f"DOC 解析失败: {exc}") from exc

        if not text or not text.strip():
            raise ValueError("DOC 文件中未提取到文本内容")

        cleaned = _clean_text(text)

        if len(cleaned) < _MIN_TEXT_LENGTH:
            warnings.append("提取的文本内容较短")

        truncated_text, truncated = _truncate_text(cleaned, max_length)

        return ParsedContent(
            text=truncated_text,
            source="doc",
            filename=filename,
            char_count=len(truncated_text),
            truncated=truncated,
            warnings=warnings,
        )

    @classmethod
    def _extract_doc_with_antiword(cls, path: Path) -> str:
        """使用 antiword 命令行工具提取 DOC 文本."""
        import subprocess

        try:
            result = subprocess.run(
                ["antiword", str(path)],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
            )
            if result.returncode == 0:
                return result.stdout
            raise ValueError(f"antiword 返回错误: {result.stderr}")
        except FileNotFoundError:
            raise ValueError("antiword 命令行工具未安装")

    # ------------------------------------------------------------------
    # TXT 解析
    # ------------------------------------------------------------------
    @classmethod
    def _from_txt_file(cls, path: Path, max_length: int) -> ParsedContent:
        """从 TXT 文件提取文本."""
        warnings: list[str] = []

        try:
            for encoding in ["utf-8", "gbk", "gb18030", "utf-16"]:
                try:
                    text = path.read_text(encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("无法识别 TXT 文件编码")

            if not text or not text.strip():
                raise ValueError("TXT 文件为空")

            cleaned = _clean_text(text)

            if len(cleaned) < _MIN_TEXT_LENGTH:
                warnings.append("提取的文本内容较短")

            truncated_text, truncated = _truncate_text(cleaned, max_length)

            return ParsedContent(
                text=truncated_text,
                source="txt",
                filename=path.name,
                char_count=len(truncated_text),
                truncated=truncated,
                warnings=warnings,
            )

        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"TXT 解析失败: {exc}") from exc

    @classmethod
    def _from_txt_bytes(cls, data: bytes, filename: Optional[str], max_length: int) -> ParsedContent:
        """从 TXT 二进制数据提取文本."""
        warnings: list[str] = []

        try:
            for encoding in ["utf-8", "gbk", "gb18030", "utf-16"]:
                try:
                    text = data.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("无法识别 TXT 文件编码")

            if not text or not text.strip():
                raise ValueError("TXT 文件为空")

            cleaned = _clean_text(text)

            if len(cleaned) < _MIN_TEXT_LENGTH:
                warnings.append("提取的文本内容较短")

            truncated_text, truncated = _truncate_text(cleaned, max_length)

            return ParsedContent(
                text=truncated_text,
                source="txt",
                filename=filename,
                char_count=len(truncated_text),
                truncated=truncated,
                warnings=warnings,
            )

        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"TXT 解析失败: {exc}") from exc

    # ------------------------------------------------------------------
    # 格式检测辅助
    # ------------------------------------------------------------------
    @staticmethod
    def _detect_format(data: bytes, filename: Optional[str] = None) -> str:
        """根据文件头和文件名检测格式."""
        if filename:
            suffix = Path(filename).suffix.lower()
            if suffix in (".pdf", ".docx", ".doc", ".txt"):
                return suffix.lstrip(".")

        if data[:4] == b"%PDF":
            return "pdf"

        if data[:4] == b"PK\x03\x04":
            return "docx"

        if data[:4] == b"\xd0\xcf\x11\xe0":
            return "doc"

        return "txt"
